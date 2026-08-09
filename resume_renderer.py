"""
resume_renderer.py

AI JOB APPLICATION & RESUME INTELLIGENCE
----------------------------------------
Converts the validated customized_resume.json into:
    1. ATS-friendly DOCX
    2. ATS-friendly PDF

Input:
    generated_resumes/customized_resume.json

Output:
    generated_resumes/<candidate>_<role>_Resume.docx
    generated_resumes/<candidate>_<role>_Resume.pdf

IMPORTANT:
- Does NOT modify master_resume.json
- Does NOT invent resume facts
- Uses the customized JSON as the source
- Preserves protected sections supplied by the customizer/validator
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ============================================================
# PATHS
# ============================================================

PROJECT_ROOT = Path(__file__).resolve().parent

INPUT_PATH = PROJECT_ROOT / "generated_resumes" / "customized_resume.json"
OUTPUT_DIR = PROJECT_ROOT / "generated_resumes"


# ============================================================
# GENERAL HELPERS
# ============================================================

def clean_text(value: Any) -> str:
    """Convert a value into clean printable text."""

    if value is None:
        return ""

    if isinstance(value, str):
        text = value
    else:
        text = str(value)

    # Convert escaped newlines if they exist
    text = text.replace("\\n", "\n")

    # Remove markdown mailto wrapper
    text = re.sub(
        r"\[([^\]]+)\]\(mailto:[^)]+\)",
        r"\1",
        text,
        flags=re.IGNORECASE,
    )

    # Remove ordinary markdown links while keeping visible text
    text = re.sub(
        r"\[([^\]]+)\]\((https?://[^)]+)\)",
        r"\1",
        text,
    )

    # Remove accidental markdown emphasis
    text = text.replace("**", "")
    text = text.replace("__", "")

    return text.strip()


def safe_filename(value: str) -> str:
    """Make a Windows-safe filename."""

    value = clean_text(value)

    value = re.sub(r'[<>:"/\\|?*]', "", value)

    value = re.sub(r"\s+", "_", value)

    return value.strip("._") or "resume"


def first_nonempty(*values: Any) -> str:
    for value in values:
        text = clean_text(value)
        if text:
            return text

    return ""


def ensure_list(value: Any) -> List[Any]:
    if value is None:
        return []

    if isinstance(value, list):
        return value

    return [value]


# ============================================================
# JSON LOADING
# ============================================================

def load_resume(path: Path) -> Dict[str, Any]:

    if not path.exists():
        raise FileNotFoundError(
            f"Customized resume not found:\n{path}\n\n"
            "Run resume_customizer.py first."
        )

    with open(path, "r", encoding="utf-8") as file:
        data = json.load(file)

    if not isinstance(data, dict):
        raise ValueError("Customized resume JSON root must be an object.")

    return data


# ============================================================
# ROLE DETECTION
# ============================================================

def detect_role(data: Dict[str, Any]) -> str:

    possible_keys = [
        "target_role",
        "target_title",
        "role",
        "job_title",
        "target_job",
        "position",
    ]

    for key in possible_keys:
        if key in data:
            value = clean_text(data.get(key))
            if value:
                return value

    metadata = data.get("metadata", {})

    if isinstance(metadata, dict):

        for key in possible_keys:
            value = clean_text(metadata.get(key))

            if value:
                return value

    customization = data.get("customization", {})

    if isinstance(customization, dict):

        for key in possible_keys:
            value = clean_text(customization.get(key))

            if value:
                return value

    return "Resume"


# ============================================================
# PERSONAL INFORMATION
# ============================================================

def get_personal(data: Dict[str, Any]) -> Dict[str, str]:

    personal = data.get("personal", {})

    if not isinstance(personal, dict):
        personal = {}

    return {
        "name": clean_text(personal.get("name", "")),
        "location": clean_text(personal.get("location", "")),
        "phone": clean_text(personal.get("phone", "")),
        "email": clean_text(personal.get("email", "")),
        "linkedin": clean_text(personal.get("linkedin", "")),
        "github": clean_text(personal.get("github", "")),
    }


# ============================================================
# SECTION TEXT
# ============================================================

def get_section_text(
    data: Dict[str, Any],
    section_name: str,
) -> str:

    section = data.get(section_name, {})

    if isinstance(section, dict):

        text = clean_text(section.get("original_text", ""))

        if text:
            return text

        items = section.get("items", [])

        if isinstance(items, list):

            return "\n".join(
                clean_text(item)
                for item in items
                if clean_text(item)
            )

    elif isinstance(section, str):

        return clean_text(section)

    return ""


# ============================================================
# SUMMARY
# ============================================================

def get_summary(data: Dict[str, Any]) -> str:

    summary = data.get("summary", {})

    if isinstance(summary, dict):

        return first_nonempty(
            summary.get("customized_text"),
            summary.get("customized"),
            summary.get("original_text"),
            summary.get("text"),
        )

    return clean_text(summary)


# ============================================================
# SKILLS
# ============================================================

def get_skills(data: Dict[str, Any]) -> Dict[str, List[str]]:

    skills = data.get("skills", {})

    if not isinstance(skills, dict):
        return {}

    categories = skills.get("categories", {})

    if not isinstance(categories, dict):
        return {}

    result = {}

    for category, values in categories.items():

        if isinstance(values, list):

            cleaned = []

            for value in values:

                text = clean_text(value)

                if text:
                    cleaned.append(text)

            if cleaned:
                result[clean_text(category)] = cleaned

        else:

            text = clean_text(values)

            if text:
                result[clean_text(category)] = [text]

    return result


# ============================================================
# PROJECTS
# ============================================================

def get_projects(data: Dict[str, Any]) -> List[Dict[str, Any]]:

    projects = data.get("projects", [])

    if not isinstance(projects, list):
        return []

    result = []

    for project in projects:

        if not isinstance(project, dict):
            continue

        name = clean_text(project.get("name", ""))

        technologies = clean_text(
            project.get("technologies", "")
        )

        bullets = []

        # Customized wording is preferred.
        for key in [
            "customized_facts",
            "customized_bullets",
            "rewritten_facts",
            "bullets",
            "facts",
        ]:

            values = project.get(key)

            if isinstance(values, list):

                for value in values:

                    text = clean_text(value)

                    if text:
                        bullets.append(text)

                if bullets:
                    break

        # Fallback to original_text if no list exists
        if not bullets:

            original = clean_text(
                project.get("original_text", "")
            )

            if original:

                lines = original.splitlines()

                for line in lines:

                    line = line.strip()

                    if line.startswith("•"):
                        line = line.lstrip("•").strip()

                    if line:
                        bullets.append(line)

        result.append(
            {
                "name": name,
                "technologies": technologies,
                "bullets": bullets,
            }
        )

    return result


# ============================================================
# VIRTUAL JOB SIMULATIONS
# ============================================================

def get_simulations(data: Dict[str, Any]) -> List[Dict[str, Any]]:

    simulations = data.get(
        "virtual_job_simulations",
        [],
    )

    if not isinstance(simulations, list):
        return []

    result = []

    for item in simulations:

        if not isinstance(item, dict):
            continue

        name = clean_text(item.get("name", ""))

        bullets = []

        facts = item.get("facts", [])

        if isinstance(facts, list):

            for fact in facts:

                text = clean_text(fact)

                if not text:
                    continue

                # Skip duplicate title/date lines
                if text == name:
                    continue

                if re.match(
                    r"^(Remote|Pune|Mumbai|Delhi|Bangalore|Hyderabad)"
                    r".*\d{4}$",
                    text,
                    re.IGNORECASE,
                ):
                    continue

                bullets.append(text)

        result.append(
            {
                "name": name,
                "bullets": bullets,
            }
        )

    return result


# ============================================================
# EDUCATION
# ============================================================

def get_education(data: Dict[str, Any]) -> List[str]:

    education = data.get("education", {})

    if not isinstance(education, dict):
        return []

    items = education.get("items", [])

    if isinstance(items, list):

        return [
            clean_text(item)
            for item in items
            if clean_text(item)
        ]

    original = clean_text(
        education.get("original_text", "")
    )

    return [
        line.strip("• ").strip()
        for line in original.splitlines()
        if line.strip()
    ]


# ============================================================
# TRAINING
# ============================================================

def get_training(data: Dict[str, Any]) -> List[str]:

    training = data.get("training", {})

    if not isinstance(training, dict):
        return []

    items = training.get("items", [])

    if isinstance(items, list):

        return [
            clean_text(item)
            for item in items
            if clean_text(item)
        ]

    original = clean_text(
        training.get("original_text", "")
    )

    return [
        line.strip("• ").strip()
        for line in original.splitlines()
        if line.strip()
    ]


# ============================================================
# EXPERIENCE
# ============================================================

def get_experience(data: Dict[str, Any]) -> List[str]:

    experience = data.get("experience", {})

    if not isinstance(experience, dict):
        return []

    items = experience.get("items", [])

    if isinstance(items, list):

        return [
            clean_text(item)
            for item in items
            if clean_text(item)
        ]

    return []


# ============================================================
# CERTIFICATIONS
# ============================================================

def get_certifications(data: Dict[str, Any]) -> List[str]:

    certifications = data.get(
        "certifications",
        {},
    )

    if not isinstance(certifications, dict):
        return []

    items = certifications.get("items", [])

    if isinstance(items, list):

        return [
            clean_text(item)
            for item in items
            if clean_text(item)
        ]

    return []


# ============================================================
# ACHIEVEMENTS
# ============================================================

def get_achievements(data: Dict[str, Any]) -> List[str]:

    achievements = data.get(
        "achievements",
        {},
    )

    if not isinstance(achievements, dict):
        return []

    items = achievements.get("items", [])

    if isinstance(items, list):

        return [
            clean_text(item)
            for item in items
            if clean_text(item)
        ]

    return []


# ============================================================
# DOCX FORMATTING HELPERS
# ============================================================

def set_cell_shading(cell, fill: str):

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = OxmlElement("w:shd")

    shd.set(qn("w:fill"), fill)

    tc_pr.append(shd)


def set_cell_margins(
    cell,
    top: int = 60,
    start: int = 80,
    bottom: int = 60,
    end: int = 80,
):

    tc = cell._tc

    tc_pr = tc.get_or_add_tcPr()

    tc_mar = tc_pr.first_child_found_in("w:tcMar")

    if tc_mar is None:

        tc_mar = OxmlElement("w:tcMar")

        tc_pr.append(tc_mar)

    for margin, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:

        node = tc_mar.find(qn(f"w:{margin}"))

        if node is None:

            node = OxmlElement(f"w:{margin}")

            tc_mar.append(node)

        node.set(qn("w:w"), str(value))

        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph):

    p = paragraph._p

    p_pr = p.get_or_add_pPr()

    p_bdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F4E79")

    p_bdr.append(bottom)

    p_pr.append(p_bdr)


def configure_docx(document: Document):

    section = document.sections[0]

    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles

    normal = styles["Normal"]

    normal.font.name = "Arial"
    normal.font.size = Pt(9)

    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0


def add_docx_section_heading(
    document: Document,
    title: str,
):

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(title.upper())

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(
        31,
        78,
        121,
    )

    add_bottom_border(paragraph)

    return paragraph


def add_docx_bullet(
    document: Document,
    text: str,
):

    paragraph = document.add_paragraph(
        style="Normal"
    )

    paragraph.style = document.styles["Normal"]

    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(
        "• " + clean_text(text)
    )

    run.font.name = "Arial"
    run.font.size = Pt(8.8)

    return paragraph


# ============================================================
# DOCX GENERATOR
# ============================================================

def create_docx(
    data: Dict[str, Any],
    output_path: Path,
):

    document = Document()

    configure_docx(document)

    personal = get_personal(data)

    # --------------------------------------------------------
    # HEADER
    # --------------------------------------------------------

    name = personal["name"] or "Candidate"

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(name.upper())

    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(
        31,
        78,
        121,
    )

    contact_parts = [
        personal["location"],
        personal["phone"],
        personal["email"],
        personal["linkedin"],
        personal["github"],
    ]

    contact_parts = [
        clean_text(x)
        for x in contact_parts
        if clean_text(x)
    ]

    if contact_parts:

        paragraph = document.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            " | ".join(contact_parts)
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)

    availability = data.get(
        "availability",
        {},
    )

    if isinstance(availability, dict):

        availability_text = clean_text(
            availability.get(
                "original_text",
                "",
            )
        )

        if availability_text:

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph.paragraph_format.space_after = Pt(3)

            run = paragraph.add_run(
                availability_text
            )

            run.italic = True
            run.font.size = Pt(8)

    # --------------------------------------------------------
    # SUMMARY
    # --------------------------------------------------------

    summary = get_summary(data)

    if summary:

        add_docx_section_heading(
            document,
            "Professional Summary",
        )

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(summary)

        run.font.name = "Arial"
        run.font.size = Pt(8.8)

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    skills = get_skills(data)

    if skills:

        add_docx_section_heading(
            document,
            "Skills",
        )

        for category, values in skills.items():

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(1)

            label = paragraph.add_run(
                f"{category}: "
            )

            label.bold = True
            label.font.name = "Arial"
            label.font.size = Pt(8.5)

            values_text = ", ".join(values)

            run = paragraph.add_run(values_text)

            run.font.name = "Arial"
            run.font.size = Pt(8.5)

    # --------------------------------------------------------
    # PROJECTS
    # --------------------------------------------------------

    projects = get_projects(data)

    if projects:

        add_docx_section_heading(
            document,
            "Projects",
        )

        for project in projects:

            name = project["name"]

            technologies = project["technologies"]

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(name)

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9.2)

            if technologies:

                run = paragraph.add_run(
                    f"  |  {technologies}"
                )

                run.italic = True
                run.font.name = "Arial"
                run.font.size = Pt(8.2)

            for bullet in project["bullets"]:

                add_docx_bullet(
                    document,
                    bullet,
                )

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    experience = get_experience(data)

    if experience:

        add_docx_section_heading(
            document,
            "Experience",
        )

        for item in experience:

            add_docx_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # VIRTUAL JOB SIMULATIONS
    # --------------------------------------------------------

    simulations = get_simulations(data)

    if simulations:

        add_docx_section_heading(
            document,
            "Virtual Job Simulations",
        )

        for simulation in simulations:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(
                simulation["name"]
            )

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8.8)

            for bullet in simulation["bullets"]:

                add_docx_bullet(
                    document,
                    bullet,
                )

    # --------------------------------------------------------
    # TRAINING
    # --------------------------------------------------------

    training = get_training(data)

    if training:

        add_docx_section_heading(
            document,
            "Training",
        )

        for item in training:

            add_docx_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # CERTIFICATIONS
    # --------------------------------------------------------

    certifications = get_certifications(data)

    if certifications:

        add_docx_section_heading(
            document,
            "Certifications",
        )

        for item in certifications:

            add_docx_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # ACHIEVEMENTS
    # --------------------------------------------------------

    achievements = get_achievements(data)

    if achievements:

        add_docx_section_heading(
            document,
            "Achievements",
        )

        for item in achievements:

            add_docx_bullet(
                document,
                item,
            )

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    education = get_education(data)

    if education:

        add_docx_section_heading(
            document,
            "Education",
        )

        for index, item in enumerate(education):

            if index == 0:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_after = Pt(1)

                run = paragraph.add_run(item)

                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8.8)

            else:

                add_docx_bullet(
                    document,
                    item,
                )

    document.save(output_path)


# ============================================================
# PDF FONT SETUP
# ============================================================

def register_pdf_fonts():

    candidates = [
        (
            "ArialUnicode",
            r"C:\Windows\Fonts\arial.ttf",
        ),
        (
            "ArialUnicode",
            r"C:\Windows\Fonts\ARIAL.TTF",
        ),
        (
            "DejaVu",
            r"C:\Windows\Fonts\DejaVuSans.ttf",
        ),
    ]

    for font_name, font_path in candidates:

        path = Path(font_path)

        if path.exists():

            try:

                pdfmetrics.registerFont(
                    TTFont(
                        font_name,
                        str(path),
                    )
                )

                return font_name

            except Exception:
                pass

    return "Helvetica"


# ============================================================
# PDF GENERATOR
# ============================================================

def create_pdf(
    data: Dict[str, Any],
    output_path: Path,
):

    font_name = register_pdf_fonts()

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "ResumeName",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=17,
        leading=19,
        alignment=TA_CENTER,
        textColor=colors.HexColor(
            "#1F4E79"
        ),
        spaceAfter=2,
    )

    contact_style = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=7.5,
        leading=9,
        alignment=TA_CENTER,
        spaceAfter=2,
    )

    availability_style = ParagraphStyle(
        "Availability",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=7.2,
        leading=9,
        alignment=TA_CENTER,
        textColor=colors.HexColor(
            "#555555"
        ),
        spaceAfter=4,
    )

    section_style = ParagraphStyle(
        "Section",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9.5,
        leading=11,
        textColor=colors.HexColor(
            "#1F4E79"
        ),
        spaceBefore=6,
        spaceAfter=3,
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8,
        leading=10,
        spaceAfter=2,
    )

    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=-6,
        spaceAfter=1,
    )

    project_style = ParagraphStyle(
        "Project",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=8.7,
        leading=10,
        spaceBefore=3,
        spaceAfter=1,
    )

    story = []

    personal = get_personal(data)

    # --------------------------------------------------------
    # HEADER
    # --------------------------------------------------------

    story.append(
        Paragraph(
            clean_text(
                personal["name"] or "Candidate"
            ).upper(),
            name_style,
        )
    )

    contact_parts = [
        personal["location"],
        personal["phone"],
        personal["email"],
        personal["linkedin"],
        personal["github"],
    ]

    contact_parts = [
        clean_text(x)
        for x in contact_parts
        if clean_text(x)
    ]

    if contact_parts:

        story.append(
            Paragraph(
                " | ".join(contact_parts),
                contact_style,
            )
        )

    availability = data.get(
        "availability",
        {},
    )

    if isinstance(availability, dict):

        availability_text = clean_text(
            availability.get(
                "original_text",
                "",
            )
        )

        if availability_text:

            story.append(
                Paragraph(
                    availability_text,
                    availability_style,
                )
            )

    # --------------------------------------------------------
    # SECTION HELPER
    # --------------------------------------------------------

    def pdf_section(title: str):

        story.append(
            Paragraph(
                title.upper(),
                section_style,
            )
        )

        story.append(
            HRFlowable(
                width="100%",
                thickness=0.6,
                color=colors.HexColor(
                    "#1F4E79"
                ),
                spaceBefore=0,
                spaceAfter=3,
            )
        )

    # --------------------------------------------------------
    # SUMMARY
    # --------------------------------------------------------

    summary = get_summary(data)

    if summary:

        pdf_section("Professional Summary")

        story.append(
            Paragraph(
                clean_text(summary),
                body_style,
            )
        )

    # --------------------------------------------------------
    # SKILLS
    # --------------------------------------------------------

    skills = get_skills(data)

    if skills:

        pdf_section("Skills")

        for category, values in skills.items():

            text = (
                f"<b>{clean_text(category)}:</b> "
                + clean_text(
                    ", ".join(values)
                )
            )

            story.append(
                Paragraph(
                    text,
                    body_style,
                )
            )

    # --------------------------------------------------------
    # PROJECTS
    # --------------------------------------------------------

    projects = get_projects(data)

    if projects:

        pdf_section("Projects")

        for project in projects:

            title = clean_text(
                project["name"]
            )

            technologies = clean_text(
                project["technologies"]
            )

            title_text = (
                f"<b>{title}</b>"
            )

            if technologies:

                title_text += (
                    f"  |  <i>{technologies}</i>"
                )

            story.append(
                Paragraph(
                    title_text,
                    project_style,
                )
            )

            for bullet in project["bullets"]:

                story.append(
                    Paragraph(
                        f"• {clean_text(bullet)}",
                        bullet_style,
                    )
                )

    # --------------------------------------------------------
    # EXPERIENCE
    # --------------------------------------------------------

    experience = get_experience(data)

    if experience:

        pdf_section("Experience")

        for item in experience:

            story.append(
                Paragraph(
                    f"• {clean_text(item)}",
                    bullet_style,
                )
            )

    # --------------------------------------------------------
    # VIRTUAL SIMULATIONS
    # --------------------------------------------------------

    simulations = get_simulations(data)

    if simulations:

        pdf_section(
            "Virtual Job Simulations"
        )

        for simulation in simulations:

            story.append(
                Paragraph(
                    f"<b>{clean_text(simulation['name'])}</b>",
                    project_style,
                )
            )

            for bullet in simulation["bullets"]:

                story.append(
                    Paragraph(
                        f"• {clean_text(bullet)}",
                        bullet_style,
                    )
                )

    # --------------------------------------------------------
    # TRAINING
    # --------------------------------------------------------

    training = get_training(data)

    if training:

        pdf_section("Training")

        for item in training:

            story.append(
                Paragraph(
                    f"• {clean_text(item)}",
                    bullet_style,
                )
            )

    # --------------------------------------------------------
    # CERTIFICATIONS
    # --------------------------------------------------------

    certifications = get_certifications(data)

    if certifications:

        pdf_section("Certifications")

        for item in certifications:

            story.append(
                Paragraph(
                    f"• {clean_text(item)}",
                    bullet_style,
                )
            )

    # --------------------------------------------------------
    # ACHIEVEMENTS
    # --------------------------------------------------------

    achievements = get_achievements(data)

    if achievements:

        pdf_section("Achievements")

        for item in achievements:

            story.append(
                Paragraph(
                    f"• {clean_text(item)}",
                    bullet_style,
                )
            )

    # --------------------------------------------------------
    # EDUCATION
    # --------------------------------------------------------

    education = get_education(data)

    if education:

        pdf_section("Education")

        for index, item in enumerate(education):

            if index == 0:

                story.append(
                    Paragraph(
                        f"<b>{clean_text(item)}</b>",
                        body_style,
                    )
                )

            else:

                story.append(
                    Paragraph(
                        f"• {clean_text(item)}",
                        bullet_style,
                    )
                )

    document.build(story)


# ============================================================
# MAIN
# ============================================================

def main() -> int:

    parser = argparse.ArgumentParser(
        description=(
            "Render validated customized resume "
            "JSON into ATS-friendly DOCX and PDF."
        )
    )

    parser.add_argument(
        "--input",
        default=str(INPUT_PATH),
        help="Path to customized_resume.json",
    )

    parser.add_argument(
        "--output-dir",
        default=str(OUTPUT_DIR),
        help="Directory for generated resume files",
    )

    args = parser.parse_args()

    input_path = Path(args.input).resolve()

    output_dir = Path(
        args.output_dir
    ).resolve()

    output_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    print("=" * 70)
    print("AI JOB APPLICATION & RESUME INTELLIGENCE")
    print("=" * 70)

    print("RESUME RENDERER")
    print()

    print(
        f"Input:\n  {input_path}"
    )

    print()

    if not input_path.exists():

        print("ERROR: customized_resume.json not found.")

        print()
        print(
            "Run resume_customizer.py first."
        )

        return 1

    try:

        data = load_resume(
            input_path
        )

    except Exception as exc:

        print(
            f"ERROR loading resume JSON:\n{exc}"
        )

        return 1

    personal = get_personal(data)

    candidate_name = (
        personal["name"]
        or "Candidate"
    )

    role = detect_role(data)

    filename_base = (
        f"{safe_filename(candidate_name)}_"
        f"{safe_filename(role)}_Resume"
    )

    docx_path = (
        output_dir
        / f"{filename_base}.docx"
    )

    pdf_path = (
        output_dir
        / f"{filename_base}.pdf"
    )

    print()
    print("=" * 70)
    print("RENDERING")
    print("=" * 70)

    try:

        print()
        print("Creating DOCX...")

        create_docx(
            data,
            docx_path,
        )

        print(
            f"✓ DOCX created:\n  {docx_path}"
        )

        print()
        print("Creating PDF...")

        create_pdf(
            data,
            pdf_path,
        )

        print(
            f"✓ PDF created:\n  {pdf_path}"
        )

    except Exception as exc:

        print()
        print("=" * 70)
        print("RENDERING FAILED")
        print("=" * 70)

        print(
            f"\nERROR: {exc}"
        )

        return 1

    print()
    print("=" * 70)
    print("RESUME RENDERING COMPLETE")
    print("=" * 70)

    print()
    print("Candidate :", candidate_name)
    print("Target    :", role)

    print()
    print("OUTPUT FILES")
    print("-" * 70)

    print(
        f"DOCX : {docx_path}"
    )

    print(
        f"PDF  : {pdf_path}"
    )

    print()
    print("STATUS: SUCCESS")

    return 0


if __name__ == "__main__":
    raise SystemExit(
        main()
    )